import streamlit as st
from io import BytesIO
import re

# DOCX
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Parsing
from PyPDF2 import PdfReader

# NLP
from sklearn.feature_extraction.text import TfidfVectorizer

# PDF (ReportLab – Streamlit Cloud friendly)
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm

st.set_page_config(page_title="AI Resume Builder PRO", page_icon="🧾", layout="wide")

# ------------------------------- Helpers --------------------------------
ACTION_VERBS = [
    "Led","Built","Developed","Created","Implemented","Optimized","Automated","Designed",
    "Improved","Resolved","Reduced","Increased","Streamlined","Deployed","Managed","Owned",
    "Analyzed","Launched","Delivered","Collaborated","Mentored","Architected","Configured",
    "Integrated","Migrated","Refactored","Tested","Documented","Monitored"
]

SKILL_TAXONOMY = {
    "Backend": {"python","java","c#","node","express","spring","django","flask","fastapi","rest","graphql","postgres","mysql","mongodb","redis"},
    "Frontend": {"html","css","javascript","typescript","react","next.js","vue","angular","tailwind","redux"},
    "Data": {"pandas","numpy","sql","power bi","tableau","matplotlib","seaborn","scikit-learn","nlp","opencv"},
    "Cloud": {"aws","gcp","azure","docker","kubernetes","terraform","linux","ci/cd","git","github actions"},
}

def clean_list(s: str):
    return [x.strip() for x in re.split(r"[,\n;]+", s or "") if x.strip()]

def has_number(s): 
    return bool(re.search(r"\b\d+(\.\d+)?%?\b", s))

def starts_with_action(s):
    return any(s.lower().startswith(v.lower()) for v in ACTION_VERBS)

def improve_bullet(raw: str, jd_terms=None):
    if not raw: return ""
    s = raw.strip()
    if not starts_with_action(s):
        s = f"{ACTION_VERBS[0]} {s}"
    # inject a relevant JD keyword if available and missing
    if jd_terms:
        for kw in jd_terms[:3]:
            if kw.lower() not in s.lower():
                s += f" using {kw}"
                break
    if not has_number(s):
        s += " — achieved ~X% improvement / saved ~X hours / impacted ~X users"
    s = re.sub(r"\s{2,}", " ", s).rstrip(". ")
    return s

def add_heading(doc, text, size=12, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(size)
    p.alignment = align
    return p

def add_bullets(doc, items):
    for it in items:
        if it.strip():
            doc.add_paragraph(it, style="List Bullet")

def extract_text_from_pdf(file):
    try:
        r = PdfReader(file)
        return "\n".join([p.extract_text() or "" for p in r.pages])
    except Exception:
        return ""

def extract_text_from_docx(file):
    # Simple, robust extraction via python-docx
    try:
        d = Document(file)
        return "\n".join(p.text for p in d.paragraphs)
    except Exception:
        return ""

def tfidf_keywords(text, top_n=20):
    text = (text or "").strip()
    if not text: return []
    vec = TfidfVectorizer(stop_words="english", token_pattern=r"(?u)\b[a-zA-Z][a-zA-Z\-\+\.#]+\b")
    X = vec.fit_transform([text])
    scores = X.toarray()[0]
    feats = vec.get_feature_names_out()
    idx = scores.argsort()[::-1][:top_n]
    return [feats[i] for i in idx if len(feats[i]) > 2]

def group_skills(skills):
    skills_lower = {s.lower() for s in skills}
    buckets = {k: [] for k in SKILL_TAXONOMY}
    misc = []
    for s in skills_lower:
        placed = False
        for bucket, vocab in SKILL_TAXONOMY.items():
            if s in vocab:
                buckets[bucket].append(s.title())
                placed = True
                break
        if not placed:
            misc.append(s.title())
    return buckets, misc

def title_block(doc, name, title, email, phone, links, align=WD_ALIGN_PARAGRAPH.LEFT, photo=None, template="Classic"):
    # Two-column header if photo provided and template not Minimal
    if photo and template in ("Classic","Modern","Two-column"):
        table = doc.add_table(rows=1, cols=2)
        table.columns[0].width = Inches(1.5)
        cell_photo = table.cell(0,0)
        cell_text = table.cell(0,1)

        # add image
        try:
            runp = cell_photo.paragraphs[0].add_run()
            runp.add_picture(photo, width=Inches(1.2))
        except Exception:
            pass

        p = cell_text.paragraphs[0]
        r = p.add_run(name)
        r.bold = True; r.font.size = Pt(18)
        if title:
            p.add_run(f" • {title}")
        p.alignment = align

        meta = " | ".join([x for x in [email, phone, links] if x])
        if meta:
            p2 = cell_text.add_paragraph(meta)
            p2.alignment = align
        return

    # simple header
    p = doc.add_paragraph()
    r = p.add_run(name)
    r.bold = True; r.font.size = Pt(18)
    if title:
        p.add_run(f" • {title}")
    p.alignment = align
    meta = " | ".join([x for x in [email, phone, links] if x])
    if meta:
        p2 = doc.add_paragraph(meta)
        p2.alignment = align

# ---------------- PDF (ReportLab) ----------------
def build_pdf_bytes(data, grouped_skills, misc_skills, jd_terms, photo_file=None):
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=1.7*cm, rightMargin=1.7*cm,
                            topMargin=1.7*cm, bottomMargin=1.7*cm)
    styles = getSampleStyleSheet()
    title = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, spaceAfter=6)
    h3 = ParagraphStyle('H3', parent=styles['Heading3'], spaceBefore=8, spaceAfter=4)
    body = styles['BodyText']

    flow = []
    # Optional photo (small, left)
    if photo_file is not None:
        try:
            # Convert to bytes for RLImage
            pf_bytes = photo_file.read()
            photo_file.seek(0)
            img = RLImage(BytesIO(pf_bytes), width=2.8*cm, height=2.8*cm)
            flow.append(img)
            flow.append(Spacer(1, 6))
        except Exception:
            pass

    # Header
    flow.append(Paragraph(data['name'], title))
    subtitle_parts = [x for x in [data.get('title'), data.get('email'), data.get('phone'), data.get('links')] if x]
    if subtitle_parts:
        flow.append(Paragraph(" • ".join(subtitle_parts), body))
        flow.append(Spacer(1, 6))

    # Summary
    if data.get("summary"):
        flow.append(Paragraph("Summary", h3))
        flow.append(Paragraph(data["summary"], body))

    # Skills
    flow.append(Paragraph("Skills", h3))
    items = []
    for k, vals in grouped_skills.items():
        if vals:
            items.append(Paragraph(f"<b>{k}:</b> {', '.join(sorted(vals))}", body))
    if misc_skills:
        items.append(Paragraph(f"<b>Other:</b> {', '.join(sorted(misc_skills))}", body))
    if data.get("soft"):
        items.append(Paragraph(f"<b>Soft:</b> {', '.join(data['soft'])}", body))
    if items:
        flow.append(ListFlowable([ListItem(i) for i in items], bulletType='bullet'))
    flow.append(Spacer(1, 4))

    # Experience
    if data["exp"]:
        flow.append(Paragraph("Experience", h3))
        for e in data["exp"]:
            header = " | ".join([x for x in [e.get('role'), e.get('company'), e.get('loc')] if x])
            if e.get("dur"): header += f" • {e['dur']}"
            flow.append(Paragraph(f"<b>{header}</b>", body))
            blts = [improve_bullet(b, jd_terms) for b in e.get("bullets", [])]
            if blts:
                flow.append(ListFlowable([ListItem(Paragraph(b, body)) for b in blts], bulletType='bullet'))

    # Projects
    if data["projects"]:
        flow.append(Paragraph("Projects", h3))
        for p in data["projects"]:
            line = p.get("name") or "Project"
            if p.get("stack"): line += f" — {p['stack']}"
            if p.get("link"): line += f" • {p['link']}"
            flow.append(Paragraph(f"<b>{line}</b>", body))
            blts = [improve_bullet(b, jd_terms) for b in p.get("bullets", [])]
            if blts:
                flow.append(ListFlowable([ListItem(Paragraph(b, body)) for b in blts], bulletType='bullet'))

    # Education
    if data.get("edu_degree") or data.get("edu_school") or data.get("edu_grad"):
        flow.append(Paragraph("Education", h3))
        edu_line = " | ".join([x for x in [data.get("edu_degree"), data.get("edu_school")] if x])
        if data.get("edu_grad"): edu_line += f" • Graduated {data['edu_grad']}"
        flow.append(Paragraph(edu_line, body))

    # Certifications
    if data.get("certs"):
        flow.append(Paragraph("Certifications", h3))
        flow.append(ListFlowable([ListItem(Paragraph(c, body)) for c in data["certs"]], bulletType='bullet'))

    # Signature (optional)
    if data.get("signature_name"):
        flow.append(Spacer(1, 8))
        flow.append(Paragraph(f"<i>Signed: {data['signature_name']}</i>", body))

    doc.build(flow)
    buf.seek(0)
    return buf.read()

# ------------------------------- Sidebar --------------------------------
with st.sidebar:
    st.title("⚙️ Settings")
    template = st.selectbox("Template", ["Classic","Modern","Minimal","Two-column"], index=0)
    include_objective = st.checkbox("Include Objective/Summary", True)
    include_soft_skills = st.checkbox("Include Soft Skills", True)
    include_cert = st.checkbox("Include Certifications", True)
    include_projects = st.checkbox("Include Projects", True)
    include_edu = st.checkbox("Include Education", True)
    st.caption("Export as DOCX or PDF. Two-column template puts Skills on the left.")

# ------------------------------- Header ---------------------------------
st.title("🧾 AI Resume Builder — PRO")
st.write("Enter details or **import an existing resume**. Paste a Job Description to tailor keywords automatically.")

# --------------------------- Import existing resume ----------------------
with st.expander("📥 Import Existing Resume (optional)"):
    imp = st.file_uploader("Upload DOCX or PDF to pre-fill fields", type=["docx","pdf"])
    imported_text = ""
    if imp:
        if imp.type.endswith("pdf"):
            imported_text = extract_text_from_pdf(imp)
        else:
            imported_text = extract_text_from_docx(imp)
        st.text_area("Extracted text (editable)", value=imported_text, height=180)

# ------------------------------- Form -----------------------------------
with st.form("resume_form"):
    c1,c2 = st.columns(2)
    with c1:
        name = st.text_input("Full Name*", "")
        title = st.text_input("Target Title", "Software Engineer")
        email = st.text_input("Email*", "")
        phone = st.text_input("Phone", "")
        links = st.text_input("Links (GitHub | LinkedIn | Portfolio)", "github.com/you | linkedin.com/in/you")
        photo_file = st.file_uploader("Profile Photo (optional)", type=["png","jpg","jpeg"])
        signature_name = st.text_input("Signature name (optional)", "")
    with c2:
        summary = st.text_area(
            "Objective / Summary",
            (imported_text[:280] + "...") if imported_text else
            "Enthusiastic developer with hands-on experience in Python, data analysis, and web apps.",
            height=110
        )

    st.markdown("### Skills")
    tech_skills = st.text_area("All Skills (comma separated)", "Python, JavaScript, HTML, CSS, SQL, Pandas, NumPy, Streamlit, AWS, Docker, React")
    soft_skills = st.text_area("Soft Skills (comma separated)", "Communication, Leadership, Problem Solving") if include_soft_skills else ""

    st.markdown("### Experience (latest first)")
    exp = []
    for i in range(1, 3+1):
        with st.expander(f"Experience #{i}"):
            role = st.text_input(f"Role #{i}", key=f"role{i}")
            company = st.text_input(f"Company #{i}", key=f"company{i}")
            loc = st.text_input(f"Location #{i}", key=f"loc{i}")
            from_to = st.text_input(f"Duration #{i} (e.g., Jun 2023 – Present)", key=f"dur{i}")
            bullets_raw = st.text_area(f"Bullets #{i} (one per line)", key=f"bul{i}", height=110)
            if role or company or bullets_raw.strip():
                exp.append({
                    "role": role, "company": company, "loc": loc, "dur": from_to,
                    "bullets": [b.strip() for b in bullets_raw.split("\n") if b.strip()]
                })

    projects = []
    if include_projects:
        st.markdown("### Projects")
        for i in range(1, 3+1):
            with st.expander(f"Project #{i}"):
                pname = st.text_input(f"Project Name #{i}", key=f"pname{i}")
                pstack = st.text_input(f"Tech / Stack #{i}", key=f"pstack{i}", help="e.g., Python, Streamlit, OpenAI API")
                pdesc_raw = st.text_area(f"Highlights (one per line) #{i}", key=f"pdesc{i}", height=90)
                plink = st.text_input(f"Link (GitHub/Live) #{i}", key=f"plink{i}")
                if pname or pdesc_raw.strip():
                    projects.append({
                        "name": pname, "stack": pstack, "link": plink,
                        "bullets": [b.strip() for b in pdesc_raw.split("\n") if b.strip()]
                    })

    edu_degree = edu_school = edu_grad = ""
    if include_edu:
        st.markdown("### Education")
        edu_degree = st.text_input("Degree", "B.Tech in Computer Science")
        edu_school = st.text_input("College/University", "XYZ Institute of Technology")
        edu_grad = st.text_input("Graduation (Year/Month)", "2026")

    certs = []
    if include_cert:
        st.markdown("### Certifications")
        certs_raw = st.text_area("Certifications (one per line)", "AWS Certified Cloud Practitioner\nGoogle Data Analytics")
        certs = [c.strip() for c in certs_raw.split("\n") if c.strip()]

    st.markdown("### 🧠 Role-Specific Tailoring (paste JD)")
    jd_text = st.text_area("Job Description text (optional, improves bullets & skills)", height=120, placeholder="Paste the JD here...")

    submitted = st.form_submit_button("Generate Resume")

# ---------------------------- Build & Export -----------------------------
if submitted:
    if not name or not email:
        st.error("Please fill at least Name and Email.")
    else:
        # JD keywords
        jd_terms = tfidf_keywords(jd_text, top_n=20) if jd_text else []

        # skills
        skills_clean = clean_list(tech_skills)
        # Infuse JD terms into skills if not already present (simple heuristic)
        for kw in jd_terms:
            kwl = kw.lower()
            if kwl.isalpha() and kwl not in [s.lower() for s in skills_clean]:
                skills_clean.append(kw)

        grouped, misc = group_skills(skills_clean)
        soft = clean_list(soft_skills) if include_soft_skills else []

        # Prepare data dict
        data = {
            "name": name, "title": title, "email": email, "phone": phone, "links": links,
            "summary": summary if include_objective else "",
            "exp": exp, "projects": projects,
            "edu_degree": edu_degree, "edu_school": edu_school, "edu_grad": edu_grad,
            "certs": certs, "signature_name": signature_name,
            "soft": soft,
        }

        # ---------------- DOCX generation ----------------
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)

        # template styles
        align = WD_ALIGN_PARAGRAPH.LEFT
        if template == "Modern":
            align = WD_ALIGN_PARAGRAPH.CENTER

        # header
        title_block(doc, name, title, email, phone, links, align=align,
                    photo=photo_file, template=template)

        # Summary
        if include_objective and summary.strip():
            add_heading(doc, "Summary", size=12, align=WD_ALIGN_PARAGRAPH.LEFT if template!="Modern" else WD_ALIGN_PARAGRAPH.CENTER)
            doc.add_paragraph(summary.strip())

        # ------------ BOLD SKILLS BLOCK ------------
        def add_skills_block(docx_doc):
            add_heading(docx_doc, "Skills", size=12)
            # Category names bold
            for k, vals in grouped.items():
                if vals:
                    p = docx_doc.add_paragraph()
                    cat = p.add_run(f"{k}: ")
                    cat.bold = True
                    p.add_run(", ".join(sorted(vals)))
            if misc:
                p = docx_doc.add_paragraph()
                cat = p.add_run("Other: ")
                cat.bold = True
                p.add_run(", ".join(sorted(misc)))
            if soft:
                p = docx_doc.add_paragraph()
                cat = p.add_run("Soft: ")
                cat.bold = True
                p.add_run(", ".join(soft))

        if template == "Two-column":
            table = doc.add_table(rows=1, cols=2)
            left, right = table.row_cells(0)
            # left column: skills + certs + links
            lp = left.paragraphs[0]
            lr = lp.add_run("SKILLS"); lr.bold=True

            # bold categories in left column
            for k, vals in grouped.items():
                if vals:
                    p = left.add_paragraph()
                    cat = p.add_run(f"{k}: ")
                    cat.bold = True
                    p.add_run(", ".join(sorted(vals)))
            if misc:
                p = left.add_paragraph()
                cat = p.add_run("Other: ")
                cat.bold = True
                p.add_run(", ".join(sorted(misc)))
            if soft:
                p = left.add_paragraph()
                cat = p.add_run("Soft: ")
                cat.bold = True
                p.add_run(", ".join(soft))

            if certs:
                left.add_paragraph("CERTIFICATIONS")
                for c in certs: left.add_paragraph(f"- {c}")
            if links:
                left.add_paragraph("LINKS")
                left.add_paragraph(links)

            # right column: experience, projects, education
            if exp:
                rp = right.add_paragraph(); rr = rp.add_run("EXPERIENCE"); rr.bold=True
                for e in exp:
                    header = " | ".join([x for x in [e['role'], e['company'], e['loc']] if x])
                    if e.get("dur"): header += f" • {e['dur']}"
                    right.add_paragraph(header)
                    improved = [improve_bullet(b, jd_terms) for b in e["bullets"]]
                    for b in improved: right.add_paragraph(b, style="List Bullet")
            if projects:
                rp = right.add_paragraph(); rr = rp.add_run("PROJECTS"); rr.bold=True
                for p in projects:
                    line = p["name"] or "Project"
                    if p.get("stack"): line += f" — {p['stack']}"
                    if p.get("link"): line += f" • {p['link']}"
                    right.add_paragraph(line)
                    improved = [improve_bullet(b, jd_terms) for b in p["bullets"]]
                    for b in improved: right.add_paragraph(b, style="List Bullet")
            if include_edu and (edu_degree or edu_school or edu_grad):
                rp = right.add_paragraph(); rr = rp.add_run("EDUCATION"); rr.bold=True
                edu_line = " | ".join([x for x in [edu_degree, edu_school] if x])
                if edu_grad: edu_line += f" • Graduated {edu_grad}"
                right.add_paragraph(edu_line)
        else:
            # single column templates
            add_skills_block(doc)

            if exp:
                add_heading(doc, "Experience", size=12)
                for e in exp:
                    header = " | ".join([x for x in [e['role'], e['company'], e['loc']] if x])
                    if e.get("dur"): header += f" • {e['dur']}"
                    doc.add_paragraph(header)
                    improved = [improve_bullet(b, jd_terms) for b in e["bullets"]]
                    add_bullets(doc, improved)

            if projects:
                add_heading(doc, "Projects", size=12)
                for p in projects:
                    line = p["name"] or "Project"
                    if p.get("stack"): line += f" — {p['stack']}"
                    if p.get("link"): line += f" • {p['link']}"
                    doc.add_paragraph(line)
                    improved = [improve_bullet(b, jd_terms) for b in p["bullets"]]
                    add_bullets(doc, improved)

            if include_edu and (edu_degree or edu_school or edu_grad):
                add_heading(doc, "Education", size=12)
                edu_line = " | ".join([x for x in [edu_degree, edu_school] if x])
                if edu_grad:
                    edu_line += f" • Graduated {edu_grad}"
                doc.add_paragraph(edu_line)

            if certs:
                add_heading(doc, "Certifications", size=12)
                add_bullets(doc, certs)

        # Signature (no date footer)
        if signature_name:
            doc.add_paragraph().add_run(f"Signed: {signature_name}").italic = True

        # DOCX bytes
        docx_bytes = BytesIO(); doc.save(docx_bytes); docx_bytes.seek(0)

        # ---------------- PDF generation (ReportLab) ----------------
        pdf_bytes = build_pdf_bytes(data, grouped, misc, jd_terms, photo_file=photo_file)

        st.success("Your tailored resume is ready!")
        cdl, cdr = st.columns(2)
        with cdl:
            st.download_button("⬇️ Download DOCX", data=docx_bytes.getvalue(),
                               file_name=f"{name.replace(' ','_')}_Resume.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with cdr:
            st.download_button("⬇️ Download PDF", data=pdf_bytes,
                               file_name=f"{name.replace(' ','_')}_Resume.pdf",
                               mime="application/pdf")

        st.caption("Tip: replace ~X placeholders with your real metrics for best impact.")
